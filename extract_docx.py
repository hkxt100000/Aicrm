#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
提取 Word 文档内容
"""

try:
    from docx import Document
    import sys
    import os
    
    # 获取文档路径
    doc_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "企业微信智能表格API文档.docx")
    
    print(f"正在读取 Word 文档: {doc_path}")
    print("=" * 100)
    
    # 读取文档
    doc = Document(doc_path)
    
    print(f"\n📄 文档总段落数: {len(doc.paragraphs)}\n")
    print("=" * 100)
    
    # 输出所有段落
    for i, para in enumerate(doc.paragraphs, 1):
        text = para.text.strip()
        if text:  # 只输出非空段落
            # 根据样式判断是否是标题
            style = para.style.name
            if 'Heading' in style or para.runs and para.runs[0].bold:
                print(f"\n{'='*100}")
                print(f"【{text}】")
                print(f"{'='*100}\n")
            else:
                print(text)
    
    # 输出表格内容
    if doc.tables:
        print(f"\n\n{'='*100}")
        print(f"文档中的表格 (共 {len(doc.tables)} 个)")
        print(f"{'='*100}\n")
        
        for table_num, table in enumerate(doc.tables, 1):
            print(f"\n--- 表格 {table_num} ---\n")
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                print(" | ".join(row_data))
            print()
    
    print("\n" + "="*100)
    print("文档读取完成！")
    print("="*100)
    
except ImportError:
    print("错误：未安装 python-docx")
    print("请运行：pip install python-docx")
    sys.exit(1)
except Exception as e:
    print(f"错误：{e}")
    import traceback
    traceback.print_exc()
    sys.exit(1)
